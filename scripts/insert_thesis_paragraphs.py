# -*- coding: utf-8 -*-
"""向论文 docx 插入答辩修改段落"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

THESIS_PATH = r"C:\Users\ASUS\Desktop\论文\论文本文\22045210-黄嘉诚毕业论文终稿_定稿版.docx"
OUTPUT_PATH = r"C:\Users\ASUS\Desktop\论文\论文本文\22045210-黄嘉诚毕业论文终稿_定稿版_答辩修订.docx"

A_STAR_PARA = (
    "本课题A*算法在Python模块grid_env_25d.py的astar_path_fallback函数中实现，"
    "与Q-Learning共用2.5D建筑栅格及十一种扩展邻域。节点状态为栅格坐标(i,j,k)，"
    "评价函数f(n)=g(n)+1.15·h(n)，其中g(n)在欧氏步长基础上叠加垂直变化惩罚与障碍净空软惩罚，"
    "使路径尽量远离建筑顶面。若Q表rollout步数超限或陷入循环，系统以当前位置为起点运行A*，"
    "将已走路径与补全航段拼接返回。Flask接口algorithm=astar可独立调用，"
    "由Java后端RestTemplate转发至前端三算法对照页展示，作为不依赖训练数据的静态栅格回退保障。"
)

GA_PARA = (
    "遗传算法由genetic_path_on_grid函数实现，与RL、A*共享建筑高度图hmap及碰撞判定。"
    "编码采用方向序列染色体：长度chromo_len=min(420,max(36,grid_n×4))，基因取值0~7对应平面八邻域移动；"
    "解码时按基因逐步仿真，Z轴由lift_z_to_clearance依据障碍高度自动抬升。"
    "种群规模默认56，迭代110代；初始化先用A*路径转换为种子染色体，再对其约1/4规模随机变异，"
    "其余个体随机生成。适应度为栅格路径长度，未达终点时叠加80倍剩余距离惩罚；"
    "每代保留前population_size/5精英，单点交叉并12%位变异产生子代，无满意解则回退A*。"
    "Flask接口algorithm=ga供三算法对照调用，便于从航程、航点数与规划耗时等维度横向评估。"
)

RL_CHOICE_PARA = (
    "综合前述三算法对比实验及系统量化评优结果，本课题选择强化学习作为路径规划核心算法具有明确依据。"
    "从评价指标看，在相同起终点与2.5D栅格环境下，强化学习路径总长度约1000 m、航点约25个、规划耗时最短；"
    "A*与遗传算法总距离均约1500 m、航点约40个，累计转弯更少、轨迹更平滑，最小障碍距离亦较稳定。"
    "系统采用归一化加权综合评分（路径长度50%、预计飞行时间30%、计算耗时10%、航点精简度10%）进行横向评优，"
    "强化学习在距离与效率两项权重最高的指标上均占优，累计距离曲线始终低于对照算法，符合低空任务对缩短航程、"
    "降低能耗的要求。同时，Q-Learning可通过势函数塑形、走廊约束及APF斥力等奖励项融入水域巡检、道路巡检等任务语义，"
    "这是A*与遗传算法难以直接表达的。考虑到纯RL在模型未覆盖区域可能到不了终点，系统保留A*与BFS回退机制，"
    "形成“强化学习主策略+传统搜索兜底”的混合方案：日常演示与实验分析以RL的短路径、高效率为主，异常场景由A*保障可达性。"
    "因此，在兼顾创新性与工程稳定性的前提下，将强化学习作为本系统首选规划方法合理可行。"
)


def style_body_paragraph(paragraph):
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(24)
    for run in paragraph.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)


def insert_after(doc, index, text):
    ref = doc.paragraphs[index]
    new_p = OxmlElement("w:p")
    ref._p.addnext(new_p)
    new_para = Paragraph(new_p, ref._parent)
    new_para.add_run(text)
    style_body_paragraph(new_para)
    return new_para


def find_anchors(doc):
    anchors = {}
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if "使前端界面始终能展示合理的飞行路径" in t:
            anchors["astar"] = i
        if "在本系统中，遗传算法主要承担对照角色" in t:
            anchors["ga"] = i
        if "司鹏搏等[24]的多智能体深度强化学习框架" in t:
            anchors["rl_analysis_end"] = i
    return anchors


def main():
    doc = Document(THESIS_PATH)
    anchors = find_anchors(doc)
    missing = {"astar", "ga", "rl_analysis_end"} - set(anchors)
    if missing:
        raise RuntimeError(f"未找到锚点段落: {missing}")

    # 避免重复插入
    existing = "\n".join(p.text for p in doc.paragraphs)
    if "astar_path_fallback函数中实现" in existing:
        for p in doc.paragraphs:
            if "astar_path_fallback函数中实现" in p.text:
                p.clear()
                run = p.add_run(A_STAR_PARA)
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run.font.size = Pt(12)
                p.paragraph_format.first_line_indent = Pt(24)
                print("已更新A*段落。")
                break
        save_path = OUTPUT_PATH
        try:
            doc.save(THESIS_PATH)
            save_path = THESIS_PATH
        except PermissionError:
            doc.save(OUTPUT_PATH)
        print("已写入:", save_path)
        return

    insert_after(doc, anchors["rl_analysis_end"], RL_CHOICE_PARA)
    insert_after(doc, anchors["ga"], GA_PARA)
    insert_after(doc, anchors["astar"], A_STAR_PARA)

    save_path = OUTPUT_PATH
    try:
        doc.save(THESIS_PATH)
        save_path = THESIS_PATH
    except PermissionError:
        doc.save(OUTPUT_PATH)
        print("原文件被占用，已另存为:", OUTPUT_PATH)

    print("已写入:", save_path)

    doc2 = Document(THESIS_PATH)
    checks = ["astar_path_fallback", "genetic_path_on_grid", "归一化加权综合评分"]
    for key in checks:
        ok = any(key in p.text for p in doc2.paragraphs)
        print(f"校验 {key}: {'通过' if ok else '失败'}")


if __name__ == "__main__":
    main()
